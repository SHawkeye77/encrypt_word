"""
Samuel Hohenshell

Program Requirements: 
    - Python3
    - python-docx module
    - A way to work with .docx files (most commonly: Microsoft Word)
    - A fake message (FAKE_FILENAME) long enough to hold the real message
    - A template (TEMPLATE_NAME) that contains no text but has all style, 
        font, margins, etc. set.

Docx Notes:
    - "document" object: An entire document. A list of "paragraph" objects
    - "paragraph" object: A block of text seperated by an "enter" key. Contains
        a list of "run" objects.
    - "run" object: A connected string of text with the exact same style.

Note: 
    - Encription is done line by line
"""

import invis_ink_setup, vigenere
import docx
from docx.shared import RGBColor, Pt

# Heading customization
HEADING_NAME     = "Samuel Hohenshell"
HEADING_SUBTITLE = "WGH, LLC"
HEADING_DATE     = "May 17, 2020"

def main():
    
    # Getting all the user input
    REAL_FILENAME, FAKE_FILENAME, TEMPLATE, RESULTING_DOC_NAME, CIPHER_KEY = \
        invis_ink_setup.setup()
    

    # Gather all text from the fake message and make each line a list item
    fake_text = docx.Document(FAKE_FILENAME)
    fake_list = []
    for paragraph in fake_text.paragraphs:
        fake_list.append(paragraph.text)

    # Gather all text from the real message and make each line a list item
    real_text = docx.Document(REAL_FILENAME)
    real_list = []
    for paragraph_num, paragraph in enumerate(real_text.paragraphs):
        if (len(paragraph.text) != 0):  # Ignore blank lines in real message
            # Making the cipher key a repeated, long string
            CIPHER_KEY = vigenere.lengthen_key(paragraph.text, CIPHER_KEY)

            # Convert all letters in the real message to their 
            #  vigenere equivalent before adding the text to the array
            real_list.append(vigenere.encrypt(paragraph.text, CIPHER_KEY))


    # Loading template, which is guide for style, font, margins, etc. (these
    #  things should already be set manually in the template file).
    doc = docx.Document(TEMPLATE)

    # Adding the letterhead to the document
    doc.add_heading(HEADING_NAME, 0)                 # Adding title
    subtitle = doc.add_heading(HEADING_SUBTITLE, 1)  # Adding subtitle
    subtitle.alignment = 1                           # Centering the subtitle
    doc.add_heading("", 1)                           # Blank line
    doc.add_paragraph(HEADING_DATE)                  # Add heading date
    doc.add_paragraph("")                            # Blank line

    length_real = len(real_list)  # Number of lines in the real message
    count_real = 0                # Current line we are on in the real message

    # Note: We assume the fake message is long enough to hold our real message

    # Going through each line in the fake message
    for line in fake_list:
        # If we have more real message to write and we encounter an empty line
        if (count_real < length_real) and (line == ""):
            # Add a line from the real message
            paragraph = doc.add_paragraph(real_list[count_real])
            # Gathering the index of the paragraph we just added
            paragraph_index = len(doc.paragraphs) - 1
            # Get the line we just added's associated "run" object
            run = doc.paragraphs[paragraph_index].runs[0]
            # Modify the font of the run we just added to be white
            run.font.color.rgb = RGBColor(255,255,255)
            # Update the line we are on for the real message
            count_real += 1
        # Otherwise...
        else:
            # Just add line from the fake message
            paragraph = doc.add_paragraph(line)

        # Ensure no "enter"s around current paragraph that'd arouse suspicion
        set_spacing(paragraph)

    # Saving the created document with specified name (NOT as template)
    doc.save(RESULTING_DOC_NAME)

    # Printing confirmation and returning
    print("Encription complete")
    return


def set_spacing(paragraph):
    # Setting line spacing before and after given "paragraph" to zero so that 
    #  there are not abnormally long spacing between fake text paragraphs,
    #  even though there will be real text hidden between the paragraphs.
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after  = Pt(0)


if __name__ == "__main__":
    main()